import re
from typing import List, Dict, Optional
import pdfplumber
from docx import Document
import os

class SyllabusParser:
    """Parse syllabus from PDF or DOCX files - extracts content, skips self-study"""
    
    def parse_file(self, file_path: str) -> List[Dict]:
        """Parse syllabus file and extract units/topics"""
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _parse_pdf(self, file_path: str) -> List[Dict]:
        """Parse PDF syllabus using pdfplumber for better layout handling"""
        try:
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Use text-based extraction as primary method
                    # This works better for syllabi that use spacing rather than table borders
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
                    
                    # Try table extraction with custom settings for borderless tables
                    # Only use if text extraction seems minimal
                    if not text or len(text.strip()) < 100:
                        table_settings = {
                            "vertical_strategy": "text",
                            "horizontal_strategy": "text",
                            "intersection_x_tolerance": 15,
                            "intersection_y_tolerance": 5,
                        }
                        tables = page.extract_tables(table_settings)
                        if tables:
                            for table in tables:
                                for row in table:
                                    # Filter None and join
                                    clean_row = [str(cell).replace('\n', ' ').strip() if cell else '' for cell in row]
                                    clean_row = [c for c in clean_row if c and c != 'None']
                                    if clean_row:
                                        row_text = " | ".join(clean_row)
                                        full_text += row_text + "\n"
            
            return self._extract_modules_regex(full_text)
        except Exception as e:
            print(f"PDF parse error: {e}")
            return []
    
    def _parse_docx(self, file_path: str) -> List[Dict]:
        """Parse DOCX syllabus"""
        try:
            doc = Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text += "\n" + row_text
            
            return self._extract_modules_regex(full_text)
        except Exception as e:
            print(f"DOCX parse error: {e}")
            return []
    
    def _extract_modules_regex(self, text: str) -> List[Dict]:
        """
        Extract modules using regex - handles table-based syllabi
        Skips 'Guided Self-Study Topics' sections
        """
        units = []
        
        # Clean up the text - normalize whitespace 
        text = self._normalize_text(text)
        
        # Regex to find Module/Unit headers
        # Matches "Module I", "UNIT 1", "Module - 1", "UNIT-III" etc.
        # Allow leading pipes or whitespace for table rows: "| Unit 1 | Title |"
        module_header_pattern = r'^[|:\s]*(?:Module|Unit|UNIT|MODULE)\s*[\-]?\s*([IVX]+|\d+)\s*[:\|\-\.]?\s*(.*?)(?:\s*(?:\d+\+\d+|\d+)\s*(?:hours?|hrs?))?$'
        
        lines = text.split('\n')
        
        current_unit = None
        current_content_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for Module Header
            # We explicitly check for common syllabus header lines to ignore them
            if "TEXT BOOK" in line.upper() or "REFERENCE" in line.upper():
                 pass

            match = re.search(module_header_pattern, line, re.IGNORECASE)
            
            # Additional check: Headers are usually short, like "UNIT 1: INTRODUCTION"
            is_header = match and len(line) < 120
            
            if is_header:
                # Save previous unit
                if current_unit:
                    content_text = "\n".join(current_content_lines)
                    topics = self._extract_topics_from_content(content_text)
                    if topics:  # Only add if we found topics
                        units.append({
                            "unitNumber": current_unit["number"],
                            "title": current_unit["title"],
                            "topics": topics
                        })
                
                # Start new unit
                unit_num_str = match.group(1)
                # Title is match group 2
                raw_title = match.group(2).strip()
                
                # Clean title: remove leading/trailing punctuation, pipes
                title = re.sub(r'^[:\|\-\.]+|[:\|\-\.]+$', '', raw_title).strip()
                # Remove "9 Hours" etc from title if missed by regex
                title = re.sub(r'\d+\s*(?:Hours|Hrs).*$', '', title, flags=re.IGNORECASE).strip()
                # Remove extraction artifacts like " |"
                title = title.replace('|', '').strip()
                # Remove isolated trailing numbers (likely hour counts from next column)
                title = re.sub(r'\s+\d+$', '', title).strip()

                if not title:
                     title = f"Unit {unit_num_str}"
                
                current_unit = {
                    "number": self._roman_to_int(unit_num_str),
                    "title": title
                }
                current_content_lines = []
            else:
                # Append to current unit content
                if current_unit is not None:
                    current_content_lines.append(line)
        
        # Save last unit
        if current_unit:
            content_text = "\n".join(current_content_lines)
            topics = self._extract_topics_from_content(content_text)
            if topics:
                units.append({
                    "unitNumber": current_unit["number"],
                    "title": current_unit["title"],
                    "topics": topics
                })
        
        return units
    
    def _normalize_text(self, text: str) -> str:
        """Normalize text - fix line breaks and whitespace"""
        # Replace multiple spaces with single space
        text = re.sub(r'[ \t]+', ' ', text)
        lines = text.split('\n')
        return '\n'.join([line.strip() for line in lines if line.strip()])
    
    def _extract_topics_from_content(self, content: str) -> List[str]:
        """Extract topics from module content, explicitly removing Guided Self-Study"""
        
        # 1. Cleaning pre-pass
        clean_lines = []
        skip = False
        
        for line in content.split('\n'):
            line_lower = line.lower()
            
            # Start skipping on known bad sections
            # "Guided Self-Study" is the specific user complaint usually
            if any(x in line_lower for x in ['guided self', 'self-study', 'self study', 'text book', 'reference book', 'course outcome']):
                skip = True
                continue
            
            # If we hit a line that looks like a main topic line (long, no keywords), maybe stop skipping?
            # It's risky. Better to assume these sections run to the end of the unit block or are contiguous blocks at the end.
            if skip:
                 # Heuristic: If we see something that looks like "Unit X" inside content (weird), stop skipping?
                 # Or just continue skipping. Usually text books are at the end.
                 continue

            clean_lines.append(line)
            
        content = "\n".join(clean_lines)
        
        # 2. Extract Topics using delimiters
        topics = []
        
        # Split by:
        # - Newlines
        # - Standalone bullets at start of line (•, -, *, ►, ▪) - NOT dashes within text
        # - Numbered lists (1. , 2.)  
        # - Semicolons followed by space (;)
        # - Period followed by space and capital letter (end of sentence)
        # - Pipes (|) - because we joined table cells with pipes
        
        # First, split by newlines and pipes
        lines = re.split(r'[\n|]', content)
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Remove leading bullets
            line = re.sub(r'^[•\-\*►▪]\s*', '', line)
            
            # Split by numbered list prefix like "1. "
            line = re.sub(r'^\d+\.\s+', '', line)
            
            # Split by semicolons followed by space
            if '; ' in line:
                sub_parts = line.split('; ')
                for sp in sub_parts:
                    sp = sp.strip()
                    if sp:
                        # Further split by period + space + capital (sentence breaks)
                        sentences = re.split(r'\.\s+(?=[A-Z])', sp)
                        for sent in sentences:
                            cleaned = self._clean_topic(sent)
                            if self._is_valid_topic(cleaned):
                                topics.append(cleaned)
            # Split by comma + space only if the line is long
            elif ', ' in line and len(line) > 50:
                sub_parts = line.split(', ')
                for sp in sub_parts:
                    cleaned = self._clean_topic(sp)
                    if self._is_valid_topic(cleaned):
                        topics.append(cleaned)
            else:
                # Keep as single topic
                cleaned = self._clean_topic(line)
                if self._is_valid_topic(cleaned):
                    topics.append(cleaned)
                    
        return topics
    
    def _is_valid_topic(self, topic: str) -> bool:
        """Check if topic is valid content"""
        if not topic or len(topic) < 3: return False
        
        topic_lower = topic.lower()
        
        # Skip common standalone words that are likely fragments
        single_word_skips = ['and', 'or', 'the', 'of', 'in', 'to', 'for', 'with', 'from', 'by', 'at', 'on']
        if topic_lower in single_word_skips:
            return False
        
        # Filter out very short single-word topics (< 4 chars) unless they're acronyms or contain numbers
        if len(topic) < 4 and ' ' not in topic:
            # Allow if it's all caps (acronym like "JSX", "CSS", "API")
            if topic.isupper():
                return True
            # Allow if it contains numbers (like "CSS3", "HTML5", "3D")
            elif any(c.isdigit() for c in topic):
                return True
            else:
                return False
        
        skip_phrases = [
            'total hours', 'periods', 'semester', 'category', 'credit', 
            'version', 'lecture', 'tutorial', 'practical', 'exam',
            'course code', 'course title'
        ]
        
        if any(phrase in topic_lower for phrase in skip_phrases): 
            return False
        
        # Skip purely numeric or duration strings
        if re.match(r'^[\d\s\W]+$', topic): return False # "12" or "-"
        if re.match(r'^\d+\s*(?:hrs?|hours?)$', topic_lower): return False
        
        return True
    
    def _clean_topic(self, text: str) -> str:
        """Clean up a topic string"""
        # Remove parens with hours like (9) or (9 Hours)
        text = re.sub(r'\(\s*\d+\s*(?:hrs?|hours?)?\s*\)', '', text, flags=re.IGNORECASE)
        # Remove numbers at start
        text = re.sub(r'^\d+[\.\)\-]\s*', '', text)
        # Remove bullets
        text = re.sub(r'^[\W_]+', '', text)
        # Remove page numbers like "12", "12-14" at end
        text = re.sub(r'\s*\d{1,3}(?:-\d{1,3})?$', '', text)
        
        return text.strip()

    def _roman_to_int(self, s: str) -> int:
        """Convert Roman numeral or digit string to integer"""
        s = s.strip().upper()
        if s.isdigit(): return int(s)
        roman_map = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10}
        return roman_map.get(s, 1)

syllabus_parser = SyllabusParser()
