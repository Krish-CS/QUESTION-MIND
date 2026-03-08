import re
from typing import List, Dict, Optional
import pdfplumber
from docx import Document
import openpyxl
import os


class CDAPParser:
    """Parse CDAP (Course Delivery and Assessment Plan) from PDF, DOCX, or Excel files.
    
    CDAP Structure:
    - Units/Modules
    - Each unit has Part 1 and Part 2 topics
    """
    
    def parse_file(self, file_path: str) -> List[Dict]:
        """Parse CDAP file and extract units with Part 1/Part 2 topics"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _parse_pdf(self, file_path: str) -> List[Dict]:
        """Parse PDF CDAP - handles both table-based and text-based PDFs"""
        try:
            all_table_rows = []
            full_text = ""
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract tables first (preferred for structured CDAPs)
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                if row and any(cell for cell in row):
                                    all_table_rows.append(row)
                    
                    # Also extract text for fallback
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
            
            # If we have table data, use table-based parsing (similar to Excel)
            if all_table_rows:
                result = self._parse_table_rows(all_table_rows)
                if result:
                    return result
            
            # Fallback to text-based parsing
            return self._extract_cdap_structure(full_text)
        except Exception as e:
            print(f"PDF parse error: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def _parse_table_rows(self, rows: List) -> List[Dict]:
        """Parse table rows (from PDF or other sources) with same logic as Excel"""
        if not rows:
            return []
        
        units = {}
        current_unit = None
        current_unit_name = None
        
        # Find header row
        header_row_idx = 0
        for idx, row in enumerate(rows[:5]):
            row_values = [str(c).strip().upper() if c else '' for c in row]
            if any('UNIT' in v or 'PART' in v or 'TOPIC' in v for v in row_values):
                header_row_idx = idx
                break
        
        # Detect column indices from header
        col_indices = {'unit': 0, 'unit_name': 1, 'part': 3, 'topic': 4, 'subtopic': 5}
        if header_row_idx < len(rows):
            header = rows[header_row_idx]
            for idx, cell in enumerate(header):
                if cell:
                    cell_upper = str(cell).upper().replace('\n', ' ')
                    if 'UNIT' in cell_upper and 'NAME' not in cell_upper and idx < 2:
                        col_indices['unit'] = idx
                    elif 'SYLLABUS' in cell_upper or ('UNIT' in cell_upper and 'NAME' in cell_upper):
                        col_indices['unit_name'] = idx
                    elif 'PART' in cell_upper:
                        col_indices['part'] = idx
                    elif 'TOPIC' in cell_upper and 'SUB' not in cell_upper:
                        col_indices['topic'] = idx
                    elif 'SUB' in cell_upper and 'TOPIC' in cell_upper:
                        col_indices['subtopic'] = idx
        
        # Parse data rows
        for row in rows[header_row_idx + 1:]:
            if not row or all(cell is None for cell in row):
                continue
            
            row_values = [str(c).strip().replace('\n', ' ') if c is not None else '' for c in row]
            
            # Skip rows with #REF! or invalid data
            if row_values and '#REF!' in row_values[0]:
                continue
            
            # Initialize variables for this row
            part_num = None
            
            # Check column for Unit Number
            unit_col = col_indices['unit']
            if len(row_values) > unit_col and row_values[unit_col]:
                unit_val = row_values[unit_col]
                if unit_val.upper() not in ['UNIT', 'UNIT NO', 'UNIT NO.', 'UNIT NUMBER', 'NONE', '']:
                    u_match = re.search(r'(?:unit|module)?\s*[:\-]?\s*([\d.]+)', unit_val, re.IGNORECASE)
                    if u_match:
                        unit_num = int(float(u_match.group(1)))
                        current_unit = unit_num
                        unit_name_col = col_indices['unit_name']
                        if len(row_values) > unit_name_col and row_values[unit_name_col]:
                            current_unit_name = row_values[unit_name_col]
            
            if current_unit is None:
                current_unit = 1
            
            # Get Part Number
            part_col = col_indices['part']
            if len(row_values) > part_col and row_values[part_col]:
                part_val = row_values[part_col]
                if part_val.upper() not in ['PART NO', 'PART NO.', 'PART NUMBER', 'PART', 'NONE', '']:
                    if part_val in ['1', '1.0', 'Part 1', 'Part-1', 'I', 'PART I']:
                        part_num = 1
                    elif part_val in ['2', '2.0', 'Part 2', 'Part-2', 'II', 'PART II']:
                        part_num = 2
                    else:
                        p_match = re.search(r'[\d.]+', part_val)
                        if p_match:
                            part_num = int(float(p_match.group()))
            
            # Get Topic
            topic_col = col_indices['topic']
            topics_to_add = []
            if len(row_values) > topic_col and row_values[topic_col]:
                topic_val = row_values[topic_col]
                if 'TOPIC' not in topic_val.upper() or len(topic_val) > 30:
                    if len(topic_val) > 3 and topic_val.upper() != 'NONE':
                        topics_to_add.append(topic_val)
            
            # Initialize unit
            if current_unit not in units:
                unit_name = current_unit_name if current_unit_name else f"Unit {current_unit}"
                units[current_unit] = {
                    "unit_number": current_unit,
                    "unit_name": unit_name,
                    "part1_topics": [],  # List of {"topic": str, "subtopics": [str]}
                    "part2_topics": []   # List of {"topic": str, "subtopics": [str]}
                }
            elif current_unit_name and units[current_unit]["unit_name"] == f"Unit {current_unit}":
                units[current_unit]["unit_name"] = current_unit_name
            
            # Add topics (with subtopics)
            for topic_text in topics_to_add:
                if topic_text:
                    # Get subtopic if available
                    subtopics = []
                    subtopic_col = col_indices.get('subtopic', 5)
                    if len(row_values) > subtopic_col and row_values[subtopic_col]:
                        subtopic_val = row_values[subtopic_col].strip()
                        if subtopic_val and len(subtopic_val) > 3 and 'SUB' not in subtopic_val.upper()[:10]:
                            subtopics.append(subtopic_val)
                    
                    target_part = part_num if part_num else 1
                    target_list = units[current_unit]["part1_topics"] if target_part == 1 else units[current_unit]["part2_topics"]
                    
                    # Check if topic already exists
                    existing_topic = next((t for t in target_list if t["topic"] == topic_text), None)
                    if existing_topic:
                        for st in subtopics:
                            if st not in existing_topic["subtopics"]:
                                existing_topic["subtopics"].append(st)
                    else:
                        target_list.append({"topic": topic_text, "subtopics": subtopics})
        
        return list(units.values())
    
    def _parse_docx(self, file_path: str) -> List[Dict]:
        """Parse DOCX CDAP"""
        try:
            doc = Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text += "\n" + row_text
            
            return self._extract_cdap_structure(full_text)
        except Exception as e:
            print(f"DOCX parse error: {e}")
            return []
    
    def _parse_excel(self, file_path: str) -> List[Dict]:
        """Parse Excel CDAP - expects columns: Unit, Part, Topic or similar
        
        Expected structure (e.g., OS CDAP.xlsx):
        - Column 0 (A): UNIT number (1.0, 2.0, etc.)
        - Column 1 (B): SYLLABUS (UNIT NAME)
        - Column 2 (C): OUTCOME (CO)
        - Column 3 (D): PART NO. (1.0 or 2.0)
        - Column 4 (E): TOPICS TO BE COVERED
        - Column 5 (F): SUB TOPICS
        - Column 6 (G): BT LEVEL
        """
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            
            units = {}
            current_unit = None
            current_unit_name = None
            
            # Find the header row (skip rows with "#REF!" or empty)
            header_row_idx = 1
            for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=5, values_only=True), start=1):
                row_values = [str(c).strip().upper() if c else '' for c in row]
                # Look for 'UNIT' or 'PART' in header
                if any('UNIT' in v or 'PART' in v or 'TOPIC' in v for v in row_values):
                    header_row_idx = row_idx
                    break
            
            # Detect column indices from header
            col_indices = {'unit': 0, 'unit_name': 1, 'part': 3, 'topic': 4, 'subtopic': 5}
            header = list(ws.iter_rows(min_row=header_row_idx, max_row=header_row_idx, values_only=True))[0]
            for idx, cell in enumerate(header):
                if cell:
                    cell_upper = str(cell).upper()
                    if 'UNIT' in cell_upper and 'NAME' not in cell_upper and idx < 2:
                        col_indices['unit'] = idx
                    elif 'SYLLABUS' in cell_upper or ('UNIT' in cell_upper and 'NAME' in cell_upper):
                        col_indices['unit_name'] = idx
                    elif 'PART' in cell_upper:
                        col_indices['part'] = idx
                    elif 'TOPIC' in cell_upper and 'SUB' not in cell_upper:
                        col_indices['topic'] = idx
                    elif 'SUB' in cell_upper and 'TOPIC' in cell_upper:
                        col_indices['subtopic'] = idx
            
            # Parse data rows (skip header)
            for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
                if not row or all(cell is None for cell in row):
                    continue
                
                row_values = [str(c).strip() if c is not None else '' for c in row]
                
                # Skip rows with #REF! or invalid data
                if '#REF!' in row_values[0] if row_values else False:
                    continue
                
                # Initialize variables for this row
                part_num = None
                topic = None
                
                # Check column for Unit Number (col 0 by default)
                unit_col = col_indices['unit']
                if len(row_values) > unit_col and row_values[unit_col]:
                    unit_val = row_values[unit_col]
                    # Skip header values
                    if unit_val.upper() not in ['UNIT', 'UNIT NO', 'UNIT NO.', 'UNIT NUMBER']:
                        # Match "Unit 1", "1", "1.0", etc.
                        u_match = re.search(r'(?:unit|module)?\s*[:\-]?\s*([\d.]+)', unit_val, re.IGNORECASE)
                        if u_match:
                            unit_num = int(float(u_match.group(1)))
                            current_unit = unit_num
                            # Get unit name from unit_name column
                            unit_name_col = col_indices['unit_name']
                            if len(row_values) > unit_name_col and row_values[unit_name_col]:
                                current_unit_name = row_values[unit_name_col]
                
                if current_unit is None:
                    current_unit = 1  # Default
                    
                # Get Part Number from the specific column (col 3 by default)
                part_col = col_indices['part']
                if len(row_values) > part_col and row_values[part_col]:
                    part_val = row_values[part_col]
                    if part_val.upper() not in ['PART NO', 'PART NO.', 'PART NUMBER', 'PART']:
                        if part_val in ['1', '1.0', 'Part 1', 'Part-1', 'I', 'PART I']:
                            part_num = 1
                        elif part_val in ['2', '2.0', 'Part 2', 'Part-2', 'II', 'PART II']:
                            part_num = 2
                        else:
                            # Try to extract number
                            p_match = re.search(r'[\d.]+', part_val)
                            if p_match:
                                part_num = int(float(p_match.group()))
                
                # Get Topic from specific column (col 4 by default) and optionally subtopic
                topic_col = col_indices['topic']
                subtopic_col = col_indices['subtopic']
                
                topics_to_add = []
                if len(row_values) > topic_col and row_values[topic_col]:
                    topic_val = row_values[topic_col]
                    # Skip header values
                    if 'TOPIC' not in topic_val.upper() or len(topic_val) > 30:
                        if len(topic_val) > 3:
                            topics_to_add.append(topic_val)
                
                # Initialize unit dict if needed
                if current_unit not in units:
                    unit_name = current_unit_name if current_unit_name else f"Unit {current_unit}"
                    units[current_unit] = {
                        "unit_number": current_unit,
                        "unit_name": unit_name,
                        "part1_topics": [],  # List of {"topic": str, "subtopics": [str]}
                        "part2_topics": []   # List of {"topic": str, "subtopics": [str]}
                    }
                elif current_unit_name and units[current_unit]["unit_name"] == f"Unit {current_unit}":
                    # Update unit name if we found a better one
                    units[current_unit]["unit_name"] = current_unit_name
                
                # Add topics to the appropriate part (with subtopics)
                for topic_text in topics_to_add:
                    if topic_text:
                        # Get subtopic if available
                        subtopics = []
                        if len(row_values) > subtopic_col and row_values[subtopic_col]:
                            subtopic_val = row_values[subtopic_col].strip()
                            if subtopic_val and len(subtopic_val) > 3 and 'SUB' not in subtopic_val.upper()[:10]:
                                subtopics.append(subtopic_val)
                        
                        # Default to Part 1 if no part number found
                        target_part = part_num if part_num else 1
                        target_list = units[current_unit]["part1_topics"] if target_part == 1 else units[current_unit]["part2_topics"]
                        
                        # Check if topic already exists, if so, add subtopic to it
                        existing_topic = next((t for t in target_list if t["topic"] == topic_text), None)
                        if existing_topic:
                            for st in subtopics:
                                if st not in existing_topic["subtopics"]:
                                    existing_topic["subtopics"].append(st)
                        else:
                            target_list.append({"topic": topic_text, "subtopics": subtopics})
            
            return list(units.values())
        except Exception as e:
            print(f"Excel parse error: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def _extract_cdap_structure(self, text: str) -> List[Dict]:
        """Extract CDAP structure from text - identifies units and Part 1/Part 2 sections"""
        units = []
        current_unit = None
        current_part = None
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect unit/module headers
            # Pattern 1: "Unit 1", "Module 1"
            unit_match = re.search(
                r'(?:unit|module)\s*[-:]?\s*([IVX\d]+)(?:\s*[-:]\s*(.+))?',
                line, re.IGNORECASE
            )
            # Pattern 2: Start of line "1 Introduction...", "2 System..." or "1 | Introduction..." (table output)
            if not unit_match:
                # Allow space OR pipe as separator
                simple_match = re.match(r'^(\d+)(?:\s+|(?:\s*[|]\s*))([A-Za-z].+)', line)
                if simple_match:
                    # Check if it looks like a unit header (short number, long title)
                    u_num = int(simple_match.group(1))
                    # Basic sanity: Unit numbers are usually 1-10
                    if 1 <= u_num <= 10: 
                        unit_match = simple_match

            if unit_match:
                unit_num = self._roman_to_int(unit_match.group(1))
                unit_name = unit_match.group(2).strip() if unit_match.group(2) else f"Unit {unit_num}"
                
                current_unit = {
                    "unit_number": unit_num,
                    "unit_name": unit_name,
                    "part1_topics": [],  # List of {"topic": str, "subtopics": [str]}
                    "part2_topics": []   # List of {"topic": str, "subtopics": [str]}
                }
                units.append(current_unit)
                continue
            
            # Detect Part 1/Part 2 headers
            part1_match = re.search(r'part\s*[-:]?\s*1|part\s*i(?:\s|$)', line, re.IGNORECASE)
            part2_match = re.search(r'part\s*[-:]?\s*2|part\s*ii(?:\s|$)', line, re.IGNORECASE)
            
            if part1_match:
                current_part = 1
                continue
            elif part2_match:
                current_part = 2
                continue
            
            # Add topics to current part
            if current_unit and current_part and self._is_valid_topic(line):
                topic = self._clean_topic(line)
                if topic:
                    topic_entry = {"topic": topic, "subtopics": []}
                    if current_part == 1:
                        if not any(t["topic"] == topic for t in current_unit["part1_topics"]):
                            current_unit["part1_topics"].append(topic_entry)
                    else:
                        if not any(t["topic"] == topic for t in current_unit["part2_topics"]):
                            current_unit["part2_topics"].append(topic_entry)
        
        return units
    
    def _is_valid_topic(self, text: str) -> bool:
        """Check if text is a valid topic"""
        text = text.strip()
        if len(text) < 5:
            return False
        if re.match(r'^[\d.]+$', text):
            return False
        if text.lower() in ['part 1', 'part 2', 'part i', 'part ii', 'topics']:
            return False
        return True
    
    def _clean_topic(self, text: str) -> str:
        """Clean topic string"""
        # Remove leading numbers/bullets
        text = re.sub(r'^[\d.•\-–]+\s*', '', text)
        # Remove extra whitespace
        text = ' '.join(text.split())
        return text.strip()
    
    def _roman_to_int(self, s: str) -> int:
        """Convert Roman numeral or digit to integer"""
        s = s.strip().upper()
        if s.isdigit():
            return int(s)
        
        roman = {'I': 1, 'V': 5, 'X': 10}
        result = 0
        prev = 0
        for char in reversed(s):
            val = roman.get(char, 0)
            if val < prev:
                result -= val
            else:
                result += val
            prev = val
        return result if result > 0 else 1


cdap_parser = CDAPParser()
